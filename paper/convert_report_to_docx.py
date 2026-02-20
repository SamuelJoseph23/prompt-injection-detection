"""
Convert simplified_report.md to simplified_report.docx
=============================================================
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

PAPER_DIR = Path(__file__).resolve().parent
PROJECT_DIR = PAPER_DIR.parent
MD_PATH = PAPER_DIR / "simplified_report.md"
DOCX_PATH = PAPER_DIR / "simplified_report.docx"

def set_margins(doc, top=1, bottom=1, left=1, right=1):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_borders(table):
    tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def parse_markdown(md_text):
    blocks = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\r")
        if line.strip().startswith("!"):
            m = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if m:
                caption = m.group(1)
                src = m.group(2)
                img_path = (PAPER_DIR / src).resolve()
                # Check for additional caption line below
                full_caption = caption
                if i + 1 < len(lines) and lines[i+1].startswith("*Figure"):
                    full_caption = lines[i+1].strip().strip("*")
                    i += 1
                blocks.append({"type": "image", "path": str(img_path), "caption": full_caption})
                i += 1
                continue
        hm = re.match(r'^(#{1,4})\s+(.+)$', line)
        if hm:
            blocks.append({"type": "heading", "level": len(hm.group(1)), "text": hm.group(2).strip()})
            i += 1
            continue
        if "|" in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].rstrip("\r")):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                stripped = lines[i].rstrip("\r").strip()
                if re.match(r'^[\s|:-]+$', stripped):
                    i += 1
                    continue
                cells = [c.strip() for c in stripped.split("|")]
                cells = [c for c in cells if c != ""]
                table_lines.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": table_lines})
            continue
        lm = re.match(r'^(\s*)([\d]+\.|\-|\*)\s+(.+)$', line)
        if lm:
            blocks.append({"type": "list_item", "indent": len(lm.group(1)), "text": lm.group(3).strip()})
            i += 1
            continue
        if line.strip():
            para_lines = [line.strip()]
            i += 1
            while i < len(lines):
                nl = lines[i].rstrip("\r")
                if not nl.strip() or nl.strip().startswith("#") or nl.strip().startswith("|") or nl.strip().startswith("!"): 
                    break
                para_lines.append(nl.strip())
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
            continue
        i += 1
    return blocks

def build_docx(blocks, output_path):
    doc = Document()
    set_margins(doc)
    add_page_numbers(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for block in blocks:
        if block["type"] == "heading":
            p = doc.add_paragraph()
            level = block["level"]
            size = 16 if level == 1 else (14 if level == 2 else 12)
            run = p.add_run(block["text"])
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif block["type"] == "paragraph":
            p = doc.add_paragraph(block["text"])
            p.paragraph_format.space_after = Pt(6)
        elif block["type"] == "list_item":
            p = doc.add_paragraph(f"• {block['text']}", style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
        elif block["type"] == "image":
            ipath = Path(block["path"])
            if ipath.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(ipath), width=Cm(14))
                if block["caption"]:
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(block["caption"])
                    cr.font.italic = True
                    cr.font.size = Pt(10)
        elif block["type"] == "table":
            rows = block["rows"]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = cell_text
                    if ri == 0:
                        set_cell_shading(cell, "D9E2F3")
            add_table_borders(table)
    doc.save(str(output_path))

if __name__ == "__main__":
    blocks = parse_markdown(MD_PATH.read_text(encoding="utf-8"))
    build_docx(blocks, DOCX_PATH)
    print(f"Saved to {DOCX_PATH}")
