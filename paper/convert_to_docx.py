"""
Convert ghost_in_the_machine.md to ghost_in_the_machine.docx
=============================================================
Produces a professionally formatted Word document with:
  - Times New Roman 12pt body, 14pt bold headings
  - 1-inch margins, 1.5 line spacing
  - Formatted tables with borders
  - Code blocks in Courier New 10pt
  - Page numbers in footer
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml


PAPER_DIR = Path(__file__).resolve().parent
MD_PATH = PAPER_DIR / "ghost_in_the_machine.md"
DOCX_PATH = PAPER_DIR / "ghost_in_the_machine.docx"


def set_margins(doc, top=1, bottom=1, left=1, right=1):
    """Set page margins in inches."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def style_paragraph(paragraph, font_name="Times New Roman", font_size=12,
                    bold=False, italic=False, alignment=None, space_after=6,
                    space_before=0, line_spacing=1.5):
    """Apply formatting to a paragraph."""
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = line_spacing
    if alignment:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic


def parse_markdown(md_text: str) -> list[dict]:
    """Parse markdown into a list of structured blocks."""
    blocks = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\r")

        # Skip horizontal rules
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # Headings
        hm = re.match(r'^(#{1,4})\s+(.+)$', line)
        if hm:
            level = len(hm.group(1))
            blocks.append({"type": "heading", "level": level, "text": hm.group(2).strip()})
            i += 1
            continue

        # Code blocks
        if line.strip().startswith("```"):
            lang = line.strip().lstrip("`").strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip("\r").strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\r"))
                i += 1
            i += 1  # skip closing ```
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].rstrip("\r")):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                stripped = lines[i].rstrip("\r").strip()
                if re.match(r'^[\s|:-]+$', stripped):
                    i += 1
                    continue
                cells = [c.strip() for c in stripped.split("|")]
                cells = [c for c in cells if c != ""]
                table_lines.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": table_lines})
            continue

        # Numbered / bulleted lists
        lm = re.match(r'^(\s*)([\d]+\.|\-|\*)\s+(.+)$', line)
        if lm:
            indent = len(lm.group(1))
            blocks.append({"type": "list_item", "indent": indent, "text": lm.group(3).strip()})
            i += 1
            continue

        # Regular paragraph text
        if line.strip():
            para_lines = [line.strip()]
            i += 1
            # Collect continuation lines (non-empty, non-special)
            while i < len(lines):
                nl = lines[i].rstrip("\r")
                if (not nl.strip() or nl.strip().startswith("#") or
                    nl.strip().startswith("```") or nl.strip().startswith("|") or
                    re.match(r'^---+$', nl.strip()) or
                    re.match(r'^(\s*)([\d]+\.|\-|\*)\s+', nl)):
                    break
                para_lines.append(nl.strip())
                i += 1
            text = " ".join(para_lines)
            blocks.append({"type": "paragraph", "text": text})
            continue

        i += 1

    return blocks


def clean_md_formatting(text: str) -> list[tuple[str, bool, bool]]:
    """Parse inline markdown (bold, italic) into (text, bold, italic) tuples."""
    segments = []
    # Simple regex-based parsing for **bold** and *italic*
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))')
    for m in pattern.finditer(text):
        if m.group(2):  # **bold**
            segments.append((m.group(2), True, False))
        elif m.group(3):  # *italic*
            segments.append((m.group(3), False, True))
        elif m.group(4):
            segments.append((m.group(4), False, False))
    if not segments:
        segments.append((text, False, False))
    return segments


def add_formatted_text(paragraph, text, font_name="Times New Roman",
                       font_size=12, base_bold=False):
    """Add text with inline markdown formatting to a paragraph."""
    segments = clean_md_formatting(text)
    for seg_text, bold, italic in segments:
        run = paragraph.add_run(seg_text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold or base_bold
        run.font.italic = italic


def build_docx(blocks: list[dict], output_path: Path):
    """Build the DOCX from parsed blocks."""
    doc = Document()

    # Page setup
    set_margins(doc)
    add_page_numbers(doc)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block["level"]
            p = doc.add_paragraph()
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_formatted_text(p, block["text"], font_size=16, base_bold=True)
                style_paragraph(p, font_size=16, bold=True,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                space_before=12, space_after=12)
            elif level == 2:
                add_formatted_text(p, block["text"], font_size=14, base_bold=True)
                style_paragraph(p, font_size=14, bold=True, space_before=12, space_after=6)
            elif level == 3:
                add_formatted_text(p, block["text"], font_size=13, base_bold=True)
                style_paragraph(p, font_size=13, bold=True, space_before=8, space_after=4)
            else:
                add_formatted_text(p, block["text"], font_size=12, base_bold=True)
                style_paragraph(p, font_size=12, bold=True, space_before=6, space_after=4)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            add_formatted_text(p, block["text"])
            fmt = p.paragraph_format
            fmt.line_spacing = 1.5
            fmt.space_after = Pt(6)

        elif btype == "list_item":
            p = doc.add_paragraph()
            indent_level = block.get("indent", 0) // 4
            p.paragraph_format.left_indent = Inches(0.5 * (indent_level + 1))
            add_formatted_text(p, f"• {block['text']}")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)

        elif btype == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.3)
            # Light gray background for code
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            p._p.get_or_add_pPr().append(shading)

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < num_cols:
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_formatted_text(p, cell_text, font_size=10,
                                           base_bold=(ri == 0))
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.line_spacing = 1.0

                        if ri == 0:
                            set_cell_shading(cell, "D9E2F3")

            add_table_borders(table)
            # Add some space after the table
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.save(str(output_path))
    print(f"DOCX saved to: {output_path}")


def main():
    print("Reading markdown...")
    md_text = MD_PATH.read_text(encoding="utf-8")

    print("Parsing markdown...")
    blocks = parse_markdown(md_text)
    print(f"  Parsed {len(blocks)} blocks")

    print("Building DOCX...")
    build_docx(blocks, DOCX_PATH)

    # Verify
    file_size = DOCX_PATH.stat().st_size
    print(f"  File size: {file_size / 1024:.1f} KB")
    print("Done!")


if __name__ == "__main__":
    main()
